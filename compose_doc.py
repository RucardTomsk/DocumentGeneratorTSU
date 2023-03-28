import os
import sys
from pathlib import Path

from docx import Document


def merge(path: str, name: str):
    master = Document()
    for f in os.listdir(path):
        if Path(f).suffix == '.docx':
            doc = Document(Path(path) / f)
            for p in doc.paragraphs:
                out_para = master.add_paragraph()
                for run in p.runs:
                    output_run = out_para.add_run(run.text)
                    output_run.bold = run.bold
                    output_run.italic = run.italic
                    output_run.underline = run.underline
                    output_run.font.color.rgb = run.font.color.rgb
                    output_run.style.name = run.style.name
                    output_run.font.name = run.font.name
                    output_run.font.size = run.font.size
        master.save("E:\Работы ТГУ\HITS\Работа\DocumentGeneratorTSU\doc/"+name)
