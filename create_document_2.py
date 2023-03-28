import string

import jinja2
import openpyxl
from docxtpl import DocxTemplate
import docx
import compose_doc


def get_start_row(max_row, sheet) -> int:
    for _row in range(1, max_row):
        if sheet.cell(row=_row, column=1).value == '+':
            return _row


def get_start_col(max_col, sheet) -> int:
    for _col in range(1, max_col):
        if sheet.cell(row=1, column=_col).value == "Курс 1":
            return _col


def get_flag_table_semestr(wb):
    sheet = wb["План"]
    if sheet.cell(row=2, column=1).value == None:
        return True
    else:
        return False


def counter_semestr(wb):
    sheet = wb["Титул"]
    value = sheet.cell(row=43, column=openpyxl.utils.column_index_from_string('C')).value
    term = int(value.split(':')[1][1])
    if get_flag_table_semestr(wb):
        return term * 2
    else:
        return term


def get_end_col(wb, max_col):
    sheet = wb["План"]
    if get_flag_table_semestr(wb):
        for _col in range(1, max_col):
            if sheet.cell(row=3, column=_col).value == "Компетенции":
                return _col
    else:
        for _col in range(1, max_col):
            if sheet.cell(row=2, column=_col).value == "Компетенции":
                return _col


def get_dict_competencies(wb):
    dict_competencies = {}
    sheet = wb["Компетенции"]
    _row = 2
    key = ''

    if sheet.max_column == 5:
        while (sheet.cell(row=_row, column=5).value != "exit"):
            if sheet.cell(row=_row, column=5).value != None and sheet.cell(row=_row, column=5).value != '-':
                col_test = 1
                while (sheet.cell(row=_row, column=col_test).value == None):
                    col_test += 1
                key = sheet.cell(row=_row, column=col_test).value
                dict_competencies[key] = [sheet.cell(row=_row, column=4).value, {}]
            else:
                if sheet.cell(row=_row, column=5).value == '-':
                    col_test = 2
                    while (sheet.cell(row=_row, column=col_test).value == None):
                        col_test += 1
                    dict_competencies[key][1][sheet.cell(row=_row, column=col_test).value] = sheet.cell(row=_row,
                                                                                                        column=4).value
            _row += 1
    else:
        while (sheet.cell(row=_row, column=6).value != "exit"):
            if sheet.cell(row=_row, column=6).value != None and sheet.cell(row=_row, column=6).value != '-':
                col_test = 1
                while (sheet.cell(row=_row, column=col_test).value == None):
                    col_test += 1
                key = sheet.cell(row=_row, column=col_test).value
                dict_competencies[key] = [sheet.cell(row=_row, column=5).value, {}]
            else:
                if sheet.cell(row=_row, column=6).value == '-':
                    col_test = 2
                    while (sheet.cell(row=_row, column=col_test).value == None):
                        col_test += 1
                    dict_competencies[key][1][sheet.cell(row=_row, column=col_test).value] = sheet.cell(row=_row,
                                                                                                        column=5).value
            _row += 1
    print(dict_competencies)
    return dict_competencies


def get_rezult_str(str_mas_cod, dict_competencies):
    mas_str = str_mas_cod.split(";")
    for index in range(len(mas_str)):
        if mas_str[index][0] == ' ':
            mas_str[index] = mas_str[index][1:len(mas_str[index])]

    str1 = ''
    str2 = ''
    dict_competencies_final = {}
    for key in dict_competencies.keys():
        for key2 in dict_competencies[key][1].keys():
            dict_competencies_final[key2] = [dict_competencies[key][1][key2], key + ' ' + dict_competencies[key][0]]
    for key in mas_str:
        str2 += key + ' ' + dict_competencies_final[key][0] + '\n\t'
        if not (dict_competencies_final[key][1] in str1):
            str1 += dict_competencies_final[key][1] + '\n\t'

    str1 = str1[:len(str1) - 2]
    str2 = str2[:len(str2) - 2]

    return str1, str2


def get_form_control(sheet, _row):
    final_str = ''
    mas_form_control = ['Экзамен', 'Зачет', 'Зачет с оценкой']
    mas = []
    for index in range(4, 7):
        all_semester = sheet.cell(row=_row, column=index).value
        if all_semester != None:
            for semester in all_semester:
                mas.append("Семестр " + semester + ', ' + mas_form_control[index - 4])

    for str_f_m in sorted(mas, key=lambda s: int(s[8])):
        final_str += str_f_m + '\n\t'

    return final_str[:len(final_str) - 2]


def get_sum_L(sheet, wb, _row, max_col, counter_semestr, d_z):
    sum_L = 0
    _col = 0
    if get_flag_table_semestr(wb):
        for col in range(1, max_col):
            if sheet.cell(row=3, column=col).value == "Лек":
                _col = col
                break
    else:
        for col in range(1, max_col):
            if sheet.cell(row=2, column=col).value == "Лек":
                _col = col
                break

    for _ in range(counter_semestr):
        value = sheet.cell(row=_row, column=_col).value
        if value != None:
            sum_L += float(value)
        _col += d_z

        if sum_L - int(sum_L) != 0:
            return str(sum_L)
        else:
            return str(int(sum_L))


def get_sum_C(sheet, wb, _row, max_col, counter_semestr, d_z):
    sum_C = 0
    _col = 0
    if get_flag_table_semestr(wb):
        for col in range(1, max_col):
            if sheet.cell(row=3, column=col).value == "Сем":
                _col = col
                break
    else:
        for col in range(1, max_col):
            if sheet.cell(row=2, column=col).value == "Сем":
                _col = col
                break
    for _ in range(counter_semestr):
        value = sheet.cell(row=_row, column=_col).value
        if value != None:
            sum_C += float(value)
        _col += d_z

    if sum_C - int(sum_C) != 0:
        return str(sum_C)
    else:
        return str(int(sum_C))


def get_sum_LR(sheet, wb, _row, max_col, counter_semestr, d_z):
    sum_LB = 0
    _col = 0
    if get_flag_table_semestr(wb):
        for col in range(1, max_col):
            if sheet.cell(row=3, column=col).value == "Лаб":
                _col = col
                break
    else:
        for col in range(1, max_col):
            if sheet.cell(row=2, column=col).value == "Лаб":
                _col = col
                break
    for _ in range(counter_semestr):
        value = sheet.cell(row=_row, column=_col).value
        if value != None:
            sum_LB += float(value)
        _col += d_z

    if sum_LB - int(sum_LB) != 0:
        return str(sum_LB)
    else:
        return str(int(sum_LB))


def get_sum_P(sheet, wb, _row, max_col, counter_semestr, d_z):
    sum_P = 0
    _col = 0
    if get_flag_table_semestr(wb):
        for col in range(1, max_col):
            if sheet.cell(row=3, column=col).value == "Пр":
                _col = col
                break
    else:
        for col in range(1, max_col):
            if sheet.cell(row=2, column=col).value == "Пр":
                _col = col
                break
    for _ in range(counter_semestr):
        value = sheet.cell(row=_row, column=_col).value
        if value != None:
            sum_P += float(value)
        _col += d_z
    if sum_P - int(sum_P) != 0:
        return str(sum_P)
    else:
        return str(int(sum_P))


def read_word_fila(name_fila, name_d):
    doc = docx.Document(name_fila)
    all_paras = doc.paragraphs
    index_start = -1
    text = ""
    for text_index in range(len(all_paras)):
        if name_d in all_paras[text_index].text:
            index_start = text_index
            break

    if index_start != -1:
        for text_index_2 in range(index_start, len(all_paras)):
            if all_paras[text_index_2].text == "Тематический план:":
                index_start = text_index_2 + 1
                break

        for text_index_3 in range(index_start, len(all_paras)):
            if not ("Тема" in all_paras[text_index_3].text):
                break
            text += '\t' + all_paras[text_index_3].text + '\n'

    if text == "":
        text = "Здесь должны быть темы"
    return text


def get_d_z(wb, start_col, max_col):
    sheet = wb["План"]
    counter = 1
    _row = 0
    if get_flag_table_semestr(wb):
        _row = 3
    else:
        _row = 2

    flag_k = sheet.cell(row=_row, column=start_col).value
    for col in range(start_col + 1, max_col):
        if sheet.cell(row=_row, column=col).value != flag_k:
            counter += 1
        else:
            break

    return counter


def get_dict_key_all_docs(name_table):
    final_dict = {}
    wb = openpyxl.load_workbook(name_table)
    dict_competencies = get_dict_competencies(wb)
    sheet = wb["Титул"]
    ss = sheet.cell(row=29, column=openpyxl.utils.column_index_from_string('D')).value
    ss_mas = ss.split(" ")
    SP = ""
    for s in range(len(ss_mas) - 1):
        SP += ss_mas[s] + " "
    final_dict["SP"] = SP[:len(SP) - 1].replace("Направление подготовки ", "")
    final_dict["PP"] = sheet.cell(row=30, column=openpyxl.utils.column_index_from_string('D')).value
    final_dict["FO"] = (sheet.cell(row=42, column=openpyxl.utils.column_index_from_string('C')).value).split(' ')[2]
    final_dict["KV"] = (sheet.cell(row=40, column=openpyxl.utils.column_index_from_string('C')).value).split(' ')[1]
    final_dict["G"] = sheet.cell(row=40, column=openpyxl.utils.column_index_from_string('W')).value

    sheet = wb["План"]
    max_rows = sheet.max_row + 1
    max_cols = sheet.max_column + 1
    col_start = get_start_col(max_cols, sheet)
    print("START_COL", col_start)
    count_semestr = counter_semestr(wb)
    print("SEMESTR", count_semestr)
    end_col = get_end_col(wb, max_cols)
    print("END_COL", end_col)
    d_z = get_d_z(wb, col_start, max_cols)

    for _row in range(get_start_row(max_rows, sheet), max_rows):
        if sheet.cell(row=_row, column=end_col - 1).value != None:
            final_dict["NAME"] = sheet.cell(row=_row, column=3).value
            print(final_dict["NAME"])
            final_dict["CODE"] = sheet.cell(row=_row, column=2).value
            final_dict["DEVELOP_GOAL"], final_dict["DEVELOP_REZULT"] = get_rezult_str(
                sheet.cell(row=_row, column=end_col).value, dict_competencies)
            final_dict["PP2"] = sheet.cell(row=_row, column=16).value
            if final_dict["PP2"] is None:
                final_dict["PPF"] = False
            else:
                final_dict["PPF"] = True
            mas_1 = final_dict["DEVELOP_GOAL"].split("\n\t")
            mas_2 = final_dict["DEVELOP_REZULT"].split("\n\t")
            mas_dict = []
            for ii in mas_1:
                a = True
                for i2 in mas_2:
                    if 'И' + ii.split()[0].replace("-", " ") + '.' in i2:
                        mas_dict.append({"K": ii, "KI": i2, "G": a})
                        a = False

            final_dict["table"] = mas_dict
            mas_places_dissiple = ['Дисциплина относится к обязательной части образовательной программы.',
                                   'Дисциплина относится к части образовательной программы, формируемой участниками образовательных отношений, является обязательной для изучения.',
                                   'Дисциплина относится к части образовательной программы, формируемой участниками образовательных отношений, предлагается обучающимся на выбор.']
            if 'О' in final_dict["CODE"]:
                final_dict["PLACE"] = mas_places_dissiple[0]
            elif ('ДВ' or 'ФТД') in final_dict["CODE"]:
                final_dict["PLACE"] = mas_places_dissiple[2]
            else:
                final_dict["PLACE"] = mas_places_dissiple[1]

            final_dict["LIGHTING_SEMES"] = get_form_control(sheet, _row)
            final_dict["ZE"] = sheet.cell(row=_row, column=openpyxl.utils.column_index_from_string('H')).value
            final_dict["K"] = sheet.cell(row=_row, column=openpyxl.utils.column_index_from_string('K')).value
            final_dict["L"] = get_sum_L(sheet, wb, _row, max_cols, count_semestr, d_z)
            if final_dict["L"] == "0":
                final_dict["LF"] = False
            else:
                final_dict["LF"] = True
            final_dict["C"] = get_sum_C(sheet, wb, _row, max_cols, count_semestr, d_z)
            if final_dict["C"] == "0":
                final_dict["CF"] = False
            else:
                final_dict["CF"] = True
            final_dict["P"] = get_sum_P(sheet, wb, _row, max_cols, count_semestr, d_z)
            if final_dict["P"] == "0":
                final_dict["PF"] = False
            else:
                final_dict["PF"] = True
            final_dict["LR"] = get_sum_LR(sheet, wb, _row, max_cols, count_semestr, d_z)
            if final_dict["LR"] == "0":
                final_dict["LRF"] = False
            else:
                final_dict["LRF"] = True
            mas_path = [".\\Resources\\Шаблоны\\Общий черновик аннотаций иностранцы.docx",
                        ".\\Resources\\Шаблоны\\Общий черновик аннотаций.docx"]
            if len(final_dict["NAME"].split("*")) > 1:
                final_dict["CONTENT"] = read_word_fila(mas_path[0], final_dict["NAME"].split("*")[0][:-1])
            else:
                final_dict["CONTENT"] = read_word_fila(mas_path[1], final_dict["NAME"])
            yield final_dict
        else:
            continue


def main():
    name_table = input("Введите полный путь к плану")
    key_plan = ["NAME", "SP", "PP", "FO", "KV", "G", "CODE", "DEVELOP_GOAL", "DEVELOP_REZULT", "PLACE",
                "LIGHTING_SEMES", "ZE", "K", "L", "C", "P", "LR", "CONTENT", "PP2", "PPF", "LF", "CF", "PF", "LRF"]
    key_FOS = ["NAME", "SP", "PP", "G", "table"]
    key_anno = ["CODE", "NAME", "PLACE", "LIGHTING_SEMES", "ZE", "K", "L", "C", "P", "LR", "CONTENT", "PP2","PPF"]
    for document in get_dict_key_all_docs(name_table):
        docPlan = DocxTemplate(".\\Resources\\Шаблоны\\Template.docx")
        contex_plan = {}
        for key in key_plan:
            contex_plan[key] = document[key]
        docPlan.render(contex_plan)
        if "*" in document["NAME"]:
            docPlan.save(
                "doc/Programs/" + document["CODE"] + " " + document["NAME"].replace('/', '_').replace('*', '_') + ".docx")
        else:
            docPlan.save("doc/Programs/" + document["CODE"] + " " + document["NAME"].replace('/', '_') + ".docx")

        docFOS = DocxTemplate("Resources/Шаблоны/shablon-fonda-otsenochnykh-sredstv-discipliny-Bak.docx")
        document["NAME"] = document["CODE"] + " " + document["NAME"]
        contex_FOS = {}
        for key in key_FOS:
            contex_FOS[key] = document[key]

        jinja_env = jinja2.Environment(autoescape=True)
        docFOS.render(contex_FOS, jinja_env)
        nameFOX = ""
        if "*" in document["NAME"]:
            nameFOX = "doc/FOS/" + document["NAME"].replace('/', '_').replace('*', '_') + ".docx"
        else:
            nameFOX = "doc/FOS/" + document["NAME"].replace('/', '_') + ".docx"

        docFOS.save(nameFOX)

        docFOX2 = docx.Document(nameFOX)
        table = docFOX2.tables[0]
        cell1 = table.cell(1, 0)
        cell2 = table.cell(1, 0)

        flag = False
        for row_index in range(2, len(table.rows)):
            if table.cell(row_index, 0).text == "":
                cell2 = table.cell(row_index, 0)
                flag = True
            else:
                if flag:
                    cell1.merge(cell2)
                    flag = False
                cell1 = table.cell(row_index, 0)

        if flag:
            cell1.merge(cell2)

        docFOX2.save(nameFOX)

        docANNO = DocxTemplate("Resources/Шаблоны/аннотации-соло.DOCX")
        contex_ANNO = {}

        if document["ZE"] is None:
            document["ZE"] = "0"

        if document["L"] == "0":
            contex_ANNO["LC"] = False
        else:
            contex_ANNO["LC"] = True

        if document["C"] == "0":
            contex_ANNO["CC"] = False
        else:
            contex_ANNO["CC"] = True

        if document["P"] == "0":
            contex_ANNO["PC"] = False
        else:
            contex_ANNO["PC"] = True

        if document["LR"] == "0":
            contex_ANNO["LRC"] = False
        else:
            contex_ANNO["LRC"] = True

        document["CONTENT"] = document["CONTENT"].replace("\t", "")
        document["LIGHTING_SEMES"] = document["LIGHTING_SEMES"].replace("\t", "")

        for key in key_anno:
            contex_ANNO[key] = document[key]

        docANNO.render(contex_ANNO)

        nameANNO = ""
        if "*" in document["NAME"]:
            nameANNO = "doc/Annotations/" + document["NAME"].replace('/', '_').replace('*', '_') + ".docx"
        else:
            nameANNO = "doc/Annotations/" + document["NAME"].replace('/', '_') + ".docx"

        docANNO.save(nameANNO)


    compose_doc.merge("E:\Работы ТГУ\HITS\Работа\DocumentGeneratorTSU\doc\Annotations", "AnnotationALL.docx")
    print("Объединение документов завершено. Документ сохранен ")
if __name__ == '__main__':
    main()
