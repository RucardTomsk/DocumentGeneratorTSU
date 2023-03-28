from docx import Document

doc = Document('Resources/Шаблоны/shablon-fonda-otsenochnykh-sredstv-discipliny-Bak.docx')

print(len(doc.paragraphs))

for text_p in doc.paragraphs:
    print(text_p.text)